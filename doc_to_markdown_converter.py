#!/usr/bin/env python3
"""
Doc to Markdown Converter for RAG Workflows
===========================================

This script converts various document formats (.doc, .docx, .html, .txt, .rtf)
to clean, well-formatted markdown files optimized for RAG (Retrieval-Augmented Generation) workflows.

Features:
- Handles multiple input formats including HTML-based exports from Confluence
- Specially handles Confluence .doc files that contain MHTML content
- Cleans up HTML content and converts to structured markdown
- Preserves document structure (headers, lists, tables, etc.)
- Removes unnecessary metadata and formatting
- Optimizes content for vector embedding and retrieval
- Batch processing support
- Configurable output formatting

Author: Claude Assistant
Date: September 17, 2025
"""

import os
import sys
import argparse
import re
import logging
import quopri
from pathlib import Path
from typing import List, Dict, Optional, Union
import tempfile
import shutil
from datetime import datetime

# Third-party imports (install with: pip install beautifulsoup4 html2text python-docx markdownify)
try:
    from bs4 import BeautifulSoup, Comment
    import html2text
    from docx import Document
    import markdownify
    import mammoth
except ImportError as e:
    print(f"Missing required packages. Please install with:")
    print("pip install beautifulsoup4 html2text python-docx markdownify mammoth")
    sys.exit(1)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('doc_converter.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


class DocToMarkdownConverter:
    """Main converter class for handling various document formats."""
    
    def __init__(self, config: Optional[Dict] = None):
        """Initialize the converter with optional configuration."""
        self.config = config or {}
        self.setup_html2text()
        
    def setup_html2text(self):
        """Configure html2text converter for optimal RAG output."""
        self.h = html2text.HTML2Text()
        self.h.ignore_links = self.config.get('ignore_links', False)
        self.h.ignore_images = self.config.get('ignore_images', True)
        self.h.ignore_emphasis = self.config.get('ignore_emphasis', False)
        self.h.body_width = self.config.get('body_width', 0)  # No line wrapping
        self.h.unicode_snob = True
        self.h.escape_snob = True
        self.h.wrap_links = False
        self.h.skip_internal_links = True
        self.h.inline_links = True
        
    def detect_file_type(self, file_path: Union[str, Path]) -> str:
        """
        Detect the file type based on content and extension.
        Special handling for Confluence .doc files that contain MHTML content.
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        # Read first few kilobytes to detect format reliably
        try:
            with open(file_path, 'rb') as f:
                header = f.read(4096)  # Read more data for better detection
                
            # Decode as text for content analysis
            try:
                header_text = header.decode('utf-8', errors='ignore')
            except:
                header_text = header.decode('latin1', errors='ignore')
            
            # Check for MHTML/MIME content first (highest priority for Confluence exports)
            if (b'MIME-Version:' in header and b'text/html' in header) or \
               ('MIME-Version:' in header_text and 'text/html' in header_text):
                logger.info(f"Detected MHTML content in {file_path} (extension: {extension})")
                return 'mhtml'
            
            # Check for HTML content
            if (b'<html' in header.lower() or b'<!doctype html' in header.lower()) or \
               ('<html' in header_text.lower() or '<!doctype html' in header_text.lower()):
                logger.info(f"Detected HTML content in {file_path}")
                return 'html'
            
            # Check for Confluence export markers (even in .doc files)
            confluence_markers = [
                'Exported From Confluence',
                'confluence.com',
                'Content-Type: multipart/related',
                'boundary="----=_Part_'
            ]
            
            if any(marker in header_text for marker in confluence_markers):
                logger.info(f"Detected Confluence export content in {file_path} (extension: {extension})")
                return 'mhtml'
            
            # Check for Word document signatures
            if header.startswith(b'PK\x03\x04'):  # ZIP signature (DOCX)
                return 'docx'
            elif header.startswith(b'\xd0\xcf\x11\xe0'):  # OLE signature (DOC)
                return 'doc'
            elif extension == '.docx':
                return 'docx'
            
            # Handle other extensions based on content and extension
            if extension in ['.rtf']:
                return 'rtf'
            elif extension in ['.txt', '.md']:
                return 'text'
            elif extension in ['.html', '.htm']:
                return 'html'
            elif extension == '.doc':
                # For .doc files that don't have binary signatures but might be text-based
                logger.warning(f"File {file_path} has .doc extension but appears to be text-based. Treating as MHTML.")
                return 'mhtml'
            else:
                logger.warning(f"Unknown file type for {file_path}, will attempt MHTML parsing")
                return 'unknown'
                
        except Exception as e:
            logger.warning(f"Could not detect file type for {file_path}: {e}")
            # Fallback based on extension
            if extension == '.docx':
                return 'docx'
            elif extension == '.doc':
                return 'mhtml'  # Assume Confluence export
            elif extension in ['.html', '.htm']:
                return 'html'
            else:
                return 'unknown'
    
    def extract_mhtml_content(self, file_path: Union[str, Path]) -> str:
        """Extract HTML content from MHTML/MIME files like Confluence exports."""
        try:
            # Try UTF-8 first, then fall back to other encodings
            encodings = ['utf-8', 'latin1', 'cp1252', 'iso-8859-1']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise ValueError("Could not decode file with any supported encoding")
        
            # Find the HTML part in the MIME message
            lines = content.split('\n')
            html_start = False
            html_content_lines = []
            in_html_part = False
            # Find the start of the HTML part
            for i, line in enumerate(lines):
                if 'Content-Type: text/html' in line:
                    in_html_part = True
                    # The actual content starts after the headers
                    content_start_index = i + 1
                    while content_start_index < len(lines) and lines[content_start_index].strip() != "":
                        content_start_index += 1
                    break

            if in_html_part:
                for i in range(content_start_index, len(lines)):
                    line = lines[i]
                    # Stop at the next MIME boundary
                    if line.startswith('------=_Part_') or line.startswith('--=='):
                        break
                    html_content_lines.append(line)

            # If no HTML section was found, try to extract HTML directly
            if not html_content_lines and ('<html' in content.lower() or '<body' in content.lower()):
                logger.info("No MIME sections found, treating entire content as HTML")
                return content

            # Join the lines and decode the entire block using quopri
            encoded_html = "".join(html_content_lines)
            # quopri.decodestring expects bytes
            decoded_html_bytes = quopri.decodestring(encoded_html.encode('latin1'))
            result = decoded_html_bytes.decode('utf-8', errors='replace')

            logger.info(f"Extracted {len(result)} characters of HTML content from MHTML")
            return result
            
        except Exception as e:
            logger.error(f"Error extracting MHTML content from {file_path}: {e}")
            # Fallback: treat the entire file as HTML
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            except Exception as e2:
                logger.error(f"Fallback reading failed: {e2}")
                raise
    
    def clean_html_content(self, html_content: str) -> str:
        """Clean and prepare HTML content for markdown conversion."""
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove unwanted elements
        unwanted_tags = ['script', 'style', 'meta', 'link', 'head']
        for tag in unwanted_tags:
            for element in soup.find_all(tag):
                element.decompose()
        
        # Remove comments
        for comment in soup.find_all(string=lambda text: isinstance(text, Comment)):
            comment.extract()
        
        # Remove or clean attributes that cause formatting issues
        for tag in soup.find_all():
            # Keep only essential attributes
            attrs_to_keep = ['href', 'src', 'alt', 'title', 'id']
            if tag.name in ['a', 'img']:
                new_attrs = {k: v for k, v in tag.attrs.items() if k in attrs_to_keep}
                tag.attrs = new_attrs
            else:
                tag.attrs = {}
        
        # Fix common Confluence/Word export issues
        # Remove empty paragraphs
        for p in soup.find_all('p'):
            if not p.get_text(strip=True):
                p.decompose()
        
        # Convert divs to paragraphs where appropriate
        for div in soup.find_all('div'):
            if not div.find_all(['div', 'p', 'ul', 'ol', 'table', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                div.name = 'p'
        
        return str(soup)
    
    def convert_docx_to_markdown(self, file_path: Union[str, Path]) -> str:
        """Convert DOCX files to markdown using mammoth."""
        try:
            with open(file_path, 'rb') as docx_file:
                result = mammoth.convert_to_html(docx_file)
                html_content = result.value
                
            # Clean and convert HTML to markdown
            cleaned_html = self.clean_html_content(html_content)
            markdown = self.h.handle(cleaned_html)
            
            return self.post_process_markdown(markdown)
            
        except Exception as e:
            logger.error(f"Error converting DOCX file {file_path}: {e}")
            # Fallback to python-docx
            return self.convert_docx_fallback(file_path)
    
    def convert_doc_to_markdown(self, file_path: Union[str, Path]) -> str:
        """
        Convert legacy DOC files to markdown.
        Note: This requires python-docx which may not work with all .doc files.
        For Confluence .doc exports that are actually MHTML, use convert_mhtml.
        """
        try:
            # First try to read as a real Word document
            doc = Document(file_path)
            markdown_content = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                    
                # Try to detect headings based on style
                if paragraph.style.name.startswith('Heading'):
                    level = paragraph.style.name.split()[-1]
                    if level.isdigit():
                        markdown_content.append(f"{'#' * int(level)} {text}")
                    else:
                        markdown_content.append(f"## {text}")
                else:
                    markdown_content.append(text)
                    
                markdown_content.append("")  # Empty line
            
            return '\n'.join(markdown_content)
            
        except Exception as e:
            logger.error(f"Could not read {file_path} as Word document: {e}")
            logger.info("File might be a Confluence export with .doc extension, trying MHTML parsing...")
            
            # Fallback to MHTML parsing for Confluence exports
            try:
                html_content = self.extract_mhtml_content(file_path)
                return self.convert_html_to_markdown(html_content)
            except Exception as e2:
                logger.error(f"MHTML fallback also failed: {e2}")
                return f"# Error Converting Document\n\nFailed to convert {file_path}\n\nErrors:\n- Word document parsing: {e}\n- MHTML parsing: {e2}"
    
    def convert_docx_fallback(self, file_path: Union[str, Path]) -> str:
        """Fallback DOCX conversion using python-docx."""
        try:
            doc = Document(file_path)
            markdown_content = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                    
                # Try to detect headings based on style
                if paragraph.style.name.startswith('Heading'):
                    level = paragraph.style.name.split()[-1]
                    if level.isdigit():
                        markdown_content.append(f"{'#' * int(level)} {text}")
                    else:
                        markdown_content.append(f"## {text}")
                else:
                    markdown_content.append(text)
                    
                markdown_content.append("")  # Empty line
            
            return '\n'.join(markdown_content)
            
        except Exception as e:
            logger.error(f"Fallback DOCX conversion failed for {file_path}: {e}")
            return f"# Error Converting Document\n\nFailed to convert {file_path}"
    
    def convert_html_to_markdown(self, html_content: str) -> str:
        """Convert HTML content to markdown."""
        cleaned_html = self.clean_html_content(html_content)
        markdown = self.h.handle(cleaned_html)
        return self.post_process_markdown(markdown)
    
    def post_process_markdown(self, markdown: str) -> str:
        """Post-process markdown for better RAG compatibility."""
        lines = markdown.split('\n')
        processed_lines = []
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            # Skip empty lines at the beginning
            if not processed_lines and not line:
                continue
            
            # Normalize headers
            if line.startswith('#'):
                # Ensure space after hash
                line = re.sub(r'^(#+)([^\s])', r'\1 \2', line)
                # Limit header depth to 6 levels
                if line.startswith('#######'):
                    line = '######' + line[7:]
            
            # Clean up list items
            if line.startswith(('- ', '* ', '+ ')):
                line = '- ' + line[2:].strip()
            elif re.match(r'^\d+\.', line):
                match = re.match(r'^(\d+)\.\s*(.*)', line)
                if match:
                    line = f"{match.group(1)}. {match.group(2).strip()}"
            
            # Remove excessive whitespace
            line = re.sub(r'\s+', ' ', line).strip()
            
            processed_lines.append(line)
        
        # Join lines and clean up excessive newlines
        result = '\n'.join(processed_lines)
        result = re.sub(r'\n{3,}', '\n\n', result)  # Max 2 consecutive newlines
        
        return result.strip()
    
    def add_metadata_header(self, content: str, original_file: Union[str, Path], 
                          title: Optional[str] = None) -> str:
        """Add metadata header for RAG optimization."""
        file_path = Path(original_file)
        
        # Extract title from content if not provided
        if not title:
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if line.startswith('# '):
                    title = line[2:].strip()
                    break
            if not title:
                title = file_path.stem.replace('_', ' ').replace('-', ' ').title()
        
        metadata = f"""---
title: "{title}"
source_file: "{file_path.name}"
converted_date: "{datetime.now().isoformat()}"
format: "markdown"
---

"""
        
        return metadata + content
    
    def convert_file(self, input_path: Union[str, Path], 
                    output_path: Optional[Union[str, Path]] = None,
                    add_metadata: bool = True) -> str:
        """Convert a single file to markdown."""
        input_path = Path(input_path)
        
        if not input_path.exists():
            raise FileNotFoundError(f"Input file not found: {input_path}")
        
        # Determine output path
        if not output_path:
            output_path = input_path.with_suffix('.md')
        else:
            output_path = Path(output_path)
        
        logger.info(f"Converting {input_path} to {output_path}")
        
        # Detect file type and convert
        file_type = self.detect_file_type(input_path)
        logger.info(f"Detected file type: {file_type}")
        
        try:
            if file_type == 'mhtml':
                html_content = self.extract_mhtml_content(input_path)
                markdown_content = self.convert_html_to_markdown(html_content)
            elif file_type == 'html':
                with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                    html_content = f.read()
                markdown_content = self.convert_html_to_markdown(html_content)
            elif file_type == 'docx':
                markdown_content = self.convert_docx_to_markdown(input_path)
            elif file_type == 'doc':
                markdown_content = self.convert_doc_to_markdown(input_path)
            elif file_type == 'text':
                with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                    markdown_content = f.read()
            else:
                # Try MHTML parsing as fallback for unknown types
                logger.info(f"Unknown file type, attempting MHTML parsing for {input_path}")
                html_content = self.extract_mhtml_content(input_path)
                markdown_content = self.convert_html_to_markdown(html_content)
            
            # Add metadata if requested
            if add_metadata:
                markdown_content = self.add_metadata_header(markdown_content, input_path)
            
            # Write output
            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            logger.info(f"Successfully converted {input_path} to {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Failed to convert {input_path}: {e}")
            raise
    
    def batch_convert(self, input_dir: Union[str, Path], 
                     output_dir: Optional[Union[str, Path]] = None,
                     pattern: str = "*.*") -> List[str]:
        """Convert multiple files in a directory."""
        input_dir = Path(input_dir)
        
        if not input_dir.exists():
            raise FileNotFoundError(f"Input directory not found: {input_dir}")
        
        if not output_dir:
            output_dir = input_dir / "markdown_output"
        else:
            output_dir = Path(output_dir)
        
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Find all matching files
        files = list(input_dir.glob(pattern))
        supported_extensions = {'.doc', '.docx', '.html', '.htm', '.txt', '.rtf'}
        
        converted_files = []
        
        for file_path in files:
            if file_path.suffix.lower() in supported_extensions or self.detect_file_type(file_path) in ['mhtml', 'html']:
                try:
                    output_path = output_dir / (file_path.stem + '.md')
                    result = self.convert_file(file_path, output_path)
                    converted_files.append(result)
                except Exception as e:
                    logger.error(f"Failed to convert {file_path}: {e}")
        
        logger.info(f"Batch conversion complete. Converted {len(converted_files)} files.")
        return converted_files


def create_requirements_txt():
    """Create requirements.txt file for the dependencies."""
    requirements = """beautifulsoup4>=4.12.0
html2text>=2020.1.16
python-docx>=0.8.11
markdownify>=0.11.6
mammoth>=1.6.0
lxml>=4.9.0
"""
    
    with open('requirements.txt', 'w') as f:
        f.write(requirements)
    
    print("Created requirements.txt file. Install dependencies with:")
    print("pip install -r requirements.txt")


def main():
    """Main function with CLI interface."""
    parser = argparse.ArgumentParser(
        description="Convert various document formats to markdown for RAG workflows",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Convert Confluence .doc export
  python doc_to_markdown_converter.py confluence_export.doc

  # Convert with custom output path
  python doc_to_markdown_converter.py input.doc -o output.md

  # Batch convert all files in directory
  python doc_to_markdown_converter.py -d documents/ -od markdown_files/

  # Convert without metadata headers
  python doc_to_markdown_converter.py input.doc --no-metadata

  # Create requirements.txt
  python doc_to_markdown_converter.py --create-requirements
        """
    )
    
    parser.add_argument('input', nargs='?', help='Input file path')
    parser.add_argument('-o', '--output', help='Output file path')
    parser.add_argument('-d', '--directory', help='Input directory for batch conversion')
    parser.add_argument('-od', '--output-directory', help='Output directory for batch conversion')
    parser.add_argument('--pattern', default='*.*', help='File pattern for batch conversion')
    parser.add_argument('--no-metadata', action='store_true', help='Skip metadata header')
    parser.add_argument('--ignore-links', action='store_true', help='Ignore links in conversion')
    parser.add_argument('--ignore-images', action='store_true', help='Ignore images in conversion')
    parser.add_argument('--create-requirements', action='store_true', help='Create requirements.txt file')
    parser.add_argument('-v', '--verbose', action='store_true', help='Verbose logging')
    
    args = parser.parse_args()
    
    if args.create_requirements:
        create_requirements_txt()
        return
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    # Configuration
    config = {
        'ignore_links': args.ignore_links,
        'ignore_images': args.ignore_images or True,  # Default to True for RAG
        'body_width': 0,  # No line wrapping for RAG
    }
    
    converter = DocToMarkdownConverter(config)
    
    try:
        if args.directory:
            # Batch conversion
            converted = converter.batch_convert(
                args.directory, 
                args.output_directory, 
                args.pattern
            )
            print(f"Converted {len(converted)} files:")
            for file_path in converted:
                print(f"  - {file_path}")
        
        elif args.input:
            # Single file conversion
            result = converter.convert_file(
                args.input, 
                args.output, 
                add_metadata=not args.no_metadata
            )
            print(f"Successfully converted to: {result}")
        
        else:
            parser.print_help()
    
    except Exception as e:
        logger.error(f"Conversion failed: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
